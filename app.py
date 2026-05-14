import streamlit as st
import cv2
import numpy as np
from PIL import Image

from src.preprocessing import full_pipeline
from src.ocr_engine import extract_text
from src.text_post_processing import fix_bullets

# NEW IMPORTS
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io

# ===============================
# PAGE CONFIG
# ===============================
st.set_page_config(
    page_title="Smart OCR Engine",
    layout="wide"
)

st.title("Smart OCR Engine")
st.caption("Extract high-quality text from images using advanced OCR and preprocessing techniques.")

# ===============================
# HELPERS (PDF + DOCX)
# ===============================

def create_docx(text):
    doc = Document()
    doc.add_heading("OCR Extracted Text", level=1)
    doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_pdf(text):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    width, height = letter
    y = height - 40

    for line in text.split("\n"):
        c.drawString(40, y, line[:100])  # simple safe wrap
        y -= 15
        if y < 40:
            c.showPage()
            y = height - 40

    c.save()
    buffer.seek(0)
    return buffer


# ===============================
# UPLOAD SECTION
# ===============================
uploaded_file = st.file_uploader(
    "Upload Image (PNG / JPG / JPEG)",
    type=["png", "jpg", "jpeg"]
)

if uploaded_file is not None:

    image = Image.open(uploaded_file).convert("RGB")
    img = np.array(image)
    img_bgr = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Original Image")
        st.image(image, use_column_width=True)

    with st.spinner("Extracting text... please wait"):
        text1 = extract_text(img_bgr)
        processed = full_pipeline(img_bgr)
        text2 = extract_text(processed)

    final_text = text1 if len(text1.strip()) > len(text2.strip()) else text2
    final_text = fix_bullets(final_text)

    with col2:
        st.subheader("Extracted Text")

        if final_text.strip():
            st.success("Text extracted successfully")
        else:
            st.warning("No strong text detected")

        st.text_area("", final_text, height=350)

    # ===============================
    # DOWNLOAD SECTION (UPDATED)
    # ===============================

    st.download_button(
        label="Download as TXT",
        data=final_text,
        file_name="ocr_output.txt",
        mime="text/plain"
    )

    pdf_file = create_pdf(final_text)
    st.download_button(
        label="Download as PDF",
        data=pdf_file,
        file_name="ocr_output.pdf",
        mime="application/pdf"
    )

    docx_file = create_docx(final_text)
    st.download_button(
        label="Download as DOCX",
        data=docx_file,
        file_name="ocr_output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # ===============================
    # DEBUG SECTION
    # ===============================
    with st.expander("Debug Info"):
        st.write("Raw OCR 1:")
        st.code(text1)

        st.write("Raw OCR 2 (Processed Image):")
        st.code(text2)